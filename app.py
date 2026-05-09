# =====================================================
# 0. IMPORTACIONES Y CONFIGURACIÓN GLOBAL
# =====================================================
"""
Modelo Tobit de Doble Censura — Intensidad de Adopción de PCS
Centro Nacional de Investigaciones de Café (Cenicafé)
Departamento del Huila | 422 fincas cafeteras
Autor: Juan Carlos Gómez Soto
Versión: 1.0 | Python 3.11 | Streamlit
"""

from __future__ import annotations

import io
import os
import warnings
from typing import Optional, Tuple

import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
import numpy as np
import pandas as pd
import plotly.express as px
import plotly.graph_objects as go
import scipy.stats as stats
import seaborn as sns
import statsmodels.api as sm
from scipy.stats import norm
from scipy.optimize import minimize
from sklearn.metrics import r2_score, mean_squared_error
import streamlit as st

# Word / Excel
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import openpyxl
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from openpyxl.utils import get_column_letter

warnings.filterwarnings("ignore")

# =====================================================
# 1. CONFIGURACIÓN STREAMLIT
# =====================================================

st.set_page_config(
    page_title="Modelo Tobit — Adopción PCS | Cenicafé",
    page_icon="☕",
    layout="wide",
    initial_sidebar_state="expanded",
    menu_items={
        "Get Help": "mailto:jcgomez@cenicafe.org",
        "About": "Análisis Tobit de Doble Censura — Cenicafé 2026",
    },
)

# Paleta oficial Cenicafé
VERDE_CAFETO = "#2E7D32"
CAFE_OSCURO = "#4E342E"
CREMA = "#FFF8E1"
DORADO = "#F9A825"
BLANCO = "#FFFFFF"

CSS = f"""
<style>
/* Sidebar */
[data-testid="stSidebar"] {{
    background: linear-gradient(180deg, {CAFE_OSCURO} 0%, {VERDE_CAFETO} 100%);
}}
[data-testid="stSidebar"] * {{
    color: {CREMA} !important;
}}
[data-testid="stSidebar"] .stSelectbox label,
[data-testid="stSidebar"] .stMultiSelect label {{
    color: {DORADO} !important;
    font-weight: 600;
}}

/* Tarjetas métricas */
[data-testid="metric-container"] {{
    background: {CREMA};
    border-left: 5px solid {VERDE_CAFETO};
    border-radius: 8px;
    padding: 10px 16px;
    box-shadow: 2px 2px 8px rgba(0,0,0,0.08);
}}

/* Títulos */
h1 {{ color: {VERDE_CAFETO} !important; }}
h2 {{ color: {CAFE_OSCURO} !important; }}
h3 {{ color: {VERDE_CAFETO} !important; border-bottom: 2px solid {DORADO}; padding-bottom: 4px; }}

/* Botones */
.stButton > button {{
    background-color: {VERDE_CAFETO};
    color: white;
    border-radius: 6px;
    border: none;
    font-weight: 600;
    padding: 8px 20px;
    transition: background 0.2s;
}}
.stButton > button:hover {{
    background-color: {CAFE_OSCURO};
}}

/* Tabs */
.stTabs [data-baseweb="tab-list"] {{
    background: {CREMA};
    border-radius: 8px;
    padding: 4px;
}}
.stTabs [data-baseweb="tab"] {{
    color: {CAFE_OSCURO};
    font-weight: 600;
}}
.stTabs [aria-selected="true"] {{
    background: {VERDE_CAFETO} !important;
    color: white !important;
    border-radius: 6px;
}}

/* Tablas dataframe */
.dataframe thead th {{
    background-color: {VERDE_CAFETO} !important;
    color: white !important;
}}

/* Divider */
.cenicafe-divider {{
    border: none;
    height: 3px;
    background: linear-gradient(90deg, {VERDE_CAFETO}, {DORADO}, {CAFE_OSCURO});
    border-radius: 3px;
    margin: 16px 0;
}}

/* Info box */
.info-box {{
    background: #E8F5E9;
    border-left: 4px solid {VERDE_CAFETO};
    border-radius: 6px;
    padding: 12px 16px;
    margin: 8px 0;
}}
</style>
"""
st.markdown(CSS, unsafe_allow_html=True)


# =====================================================
# 2. FUNCIONES AUXILIARES
# =====================================================

def divider() -> None:
    """Línea divisora con gradiente Cenicafé."""
    st.markdown('<hr class="cenicafe-divider">', unsafe_allow_html=True)


def info_box(text: str) -> None:
    """Caja de información estilizada."""
    st.markdown(f'<div class="info-box">ℹ️ {text}</div>', unsafe_allow_html=True)


def metric_row(items: list[dict]) -> None:
    """Fila de métricas KPI."""
    cols = st.columns(len(items))
    for col, item in zip(cols, items):
        with col:
            st.metric(
                label=item.get("label", ""),
                value=item.get("value", ""),
                delta=item.get("delta", None),
            )


def fig_to_bytes(fig: plt.Figure, dpi: int = 150) -> bytes:
    """Convierte figura matplotlib a bytes PNG."""
    buf = io.BytesIO()
    fig.savefig(buf, format="png", dpi=dpi, bbox_inches="tight", facecolor="white")
    buf.seek(0)
    return buf.getvalue()


def styled_df(df: pd.DataFrame) -> pd.DataFrame:
    """Devuelve dataframe con formateo numérico de 4 decimales."""
    return df.round(4)


# =====================================================
# 3. CARGA Y VALIDACIÓN DE DATOS
# =====================================================

DEFAULT_FILE = "TobitDataPCS_MASTER_Python.csv"
REQUIRED_COL = "Intensidad"
EXCLUDE_COLS = ["Nume", "Mpio"]
LOWER_LIMIT = 0.0
UPPER_LIMIT = 100.0


@st.cache_data(show_spinner=False)
def load_default_data() -> Optional[pd.DataFrame]:
    """Carga automática del dataset por defecto si existe."""
    if os.path.exists(DEFAULT_FILE):
        try:
            df = pd.read_csv(DEFAULT_FILE, sep=";")
            return df
        except Exception as e:
            st.warning(f"No se pudo cargar el archivo por defecto: {e}")
    return None


def validate_dataset(df: pd.DataFrame) -> Tuple[bool, list[str]]:
    """
    Valida el dataset para uso en modelo Tobit.
    Retorna (es_válido, lista_de_errores).
    """
    errors: list[str] = []
    warnings_list: list[str] = []

    # Existencia de variable dependiente
    if REQUIRED_COL not in df.columns:
        errors.append(f"❌ Columna '{REQUIRED_COL}' no encontrada.")
        return False, errors

    # Valores en rango [0, 100]
    y = pd.to_numeric(df[REQUIRED_COL], errors="coerce")
    if y.isnull().any():
        errors.append(f"❌ '{REQUIRED_COL}' contiene valores no numéricos.")
    if y.dropna().lt(LOWER_LIMIT).any() or y.dropna().gt(UPPER_LIMIT).any():
        errors.append(f"❌ '{REQUIRED_COL}' tiene valores fuera del rango [0, 100].")

    # Columnas vacías
    empty_cols = [c for c in df.columns if df[c].isnull().all()]
    if empty_cols:
        warnings_list.append(f"⚠️ Columnas completamente vacías: {empty_cols}")

    # Columnas duplicadas
    if df.columns.duplicated().any():
        errors.append("❌ Existen columnas duplicadas.")

    # Tipos de datos — detectar columnas object no excluidas
    non_numeric = [
        c for c in df.columns
        if c not in EXCLUDE_COLS and df[c].dtype == object
    ]
    if non_numeric:
        warnings_list.append(f"⚠️ Columnas no numéricas (serán excluidas): {non_numeric}")

    # Infinitos
    num_df = df.select_dtypes(include=[np.number])
    if np.isinf(num_df.values).any():
        errors.append("❌ Existen valores infinitos en el dataset.")

    # Desviación estándar cero (columnas constantes)
    zero_std = [c for c in num_df.columns if num_df[c].std() == 0 and c != REQUIRED_COL]
    if zero_std:
        warnings_list.append(f"⚠️ Columnas con varianza cero (excluidas): {zero_std}")

    for w in warnings_list:
        st.warning(w)

    is_valid = len(errors) == 0
    return is_valid, errors


# =====================================================
# 4. PREPROCESAMIENTO
# =====================================================

@st.cache_data(show_spinner=False)
def preprocess(
    df: pd.DataFrame,
    winsorize: bool = False,
    winsorize_pct: float = 0.05,
) -> Tuple[pd.DataFrame, pd.Series, list[str]]:
    """
    Preprocesa el dataset:
    - Coerción numérica
    - Eliminación de columnas inválidas
    - Winsorización opcional
    Retorna (df_limpio, y, features_usados)
    """
    df = df.copy()

    # Coerción numérica de todas las columnas excepto excluidas
    for col in df.columns:
        if col not in EXCLUDE_COLS:
            df[col] = pd.to_numeric(df[col], errors="coerce")

    # Eliminar columnas de exclusión y no numéricas
    drop_cols = [c for c in EXCLUDE_COLS if c in df.columns]
    df_model = df.drop(columns=drop_cols, errors="ignore")

    # Eliminar columnas con varianza cero o todo nulos
    zero_std_cols = [c for c in df_model.columns if df_model[c].std() == 0]
    all_null_cols = [c for c in df_model.columns if df_model[c].isnull().all()]
    df_model = df_model.drop(columns=zero_std_cols + all_null_cols, errors="ignore")

    # Eliminar filas con nulos
    df_model = df_model.dropna()

    # Variable dependiente
    y = df_model[REQUIRED_COL].copy()

    # Features
    X_cols = [c for c in df_model.columns if c != REQUIRED_COL]

    # Winsorización opcional
    if winsorize:
        for col in X_cols:
            lower = df_model[col].quantile(winsorize_pct)
            upper = df_model[col].quantile(1 - winsorize_pct)
            df_model[col] = df_model[col].clip(lower, upper)
        lo_y = y.quantile(winsorize_pct)
        hi_y = y.quantile(1 - winsorize_pct)
        y = y.clip(lo_y, hi_y)

    return df_model, y, X_cols


# =====================================================
# 5. MODELO TOBIT
# =====================================================

class DoubleCensoredTobit:
    """
    Modelo Tobit de doble censura implementado con
    estimación por Máxima Verosimilitud (MLE).

    Modelo latente: y* = Xβ + ε,  ε ~ N(0, σ²)
    Variable observada:
        y = 0    si y* ≤ 0
        y = y*   si 0 < y* < 100
        y = 100  si y* ≥ 100

    Los coeficientes β representan efectos sobre la
    VARIABLE LATENTE, NO sobre la variable observada.
    """

    def __init__(self, lower: float = 0.0, upper: float = 100.0):
        self.lower = lower
        self.upper = upper
        self.params: Optional[np.ndarray] = None
        self.se: Optional[np.ndarray] = None
        self.result = None
        self.feature_names: list[str] = []
        self.n_obs: int = 0
        self.loglik: float = np.nan
        self.sigma: float = np.nan

    def _neg_loglik(self, params: np.ndarray, X: np.ndarray, y: np.ndarray) -> float:
        """Log-verosimilitud negativa del modelo Tobit de doble censura."""
        beta = params[:-1]
        log_sigma = params[-1]
        sigma = np.exp(log_sigma)
        if sigma <= 0:
            return 1e10

        xb = X @ beta
        ll = 0.0

        # Observaciones censuradas en 0
        mask_low = y <= self.lower
        if mask_low.any():
            z_low = (self.lower - xb[mask_low]) / sigma
            prob = norm.cdf(z_low)
            prob = np.clip(prob, 1e-300, None)
            ll += np.sum(np.log(prob))

        # Observaciones censuradas en 100
        mask_high = y >= self.upper
        if mask_high.any():
            z_high = (self.upper - xb[mask_high]) / sigma
            prob = 1 - norm.cdf(z_high)
            prob = np.clip(prob, 1e-300, None)
            ll += np.sum(np.log(prob))

        # Observaciones no censuradas
        mask_mid = ~mask_low & ~mask_high
        if mask_mid.any():
            z = (y[mask_mid] - xb[mask_mid]) / sigma
            ll += np.sum(norm.logpdf(z) - np.log(sigma))

        return -ll

    def fit(self, X: pd.DataFrame, y: pd.Series) -> "DoubleCensoredTobit":
        """Ajusta el modelo Tobit por MLE."""
        self.feature_names = list(X.columns)
        self.n_obs = len(y)

        X_np = X.values.astype(float)
        y_np = y.values.astype(float)

        # Valores iniciales: OLS
        try:
            ols = sm.OLS(y_np, X_np).fit()
            beta0 = ols.params
        except Exception:
            beta0 = np.zeros(X_np.shape[1])
        log_sigma0 = np.log(np.std(y_np - X_np @ beta0) + 1e-8)
        params0 = np.append(beta0, log_sigma0)

        # Optimización
        res = minimize(
            self._neg_loglik,
            params0,
            args=(X_np, y_np),
            method="L-BFGS-B",
            options={"maxiter": 2000, "ftol": 1e-12, "gtol": 1e-8},
        )
        self.result = res
        self.params = res.x
        self.sigma = np.exp(res.x[-1])
        self.loglik = -res.fun

        # Errores estándar por Hessiana numérica
        try:
            from statsmodels.tools.numdiff import approx_hess
            hess = approx_hess(res.x, self._neg_loglik, args=(X_np, y_np))
            cov = np.linalg.inv(hess)
            self.se = np.sqrt(np.diag(np.abs(cov)))
        except Exception:
            self.se = np.full_like(self.params, np.nan)

        return self

    def summary_df(self) -> pd.DataFrame:
        """Tabla de resultados del modelo."""
        if self.params is None:
            raise ValueError("Modelo no ajustado.")

        beta = self.params[:-1]
        se = self.se[:-1] if self.se is not None else np.full_like(beta, np.nan)
        z_vals = beta / (se + 1e-300)
        p_vals = 2 * (1 - norm.cdf(np.abs(z_vals)))
        ci_low = beta - 1.96 * se
        ci_high = beta + 1.96 * se

        stars = []
        for p in p_vals:
            if p < 0.01:
                stars.append("***")
            elif p < 0.05:
                stars.append("**")
            elif p < 0.10:
                stars.append("*")
            else:
                stars.append("")

        df = pd.DataFrame(
            {
                "Variable": self.feature_names,
                "Coeficiente": beta,
                "Error Estándar": se,
                "z-Valor": z_vals,
                "p-Valor": p_vals,
                "Sig.": stars,
                "IC 95% Inf.": ci_low,
                "IC 95% Sup.": ci_high,
            }
        )
        return df.round(4)

    def fit_stats(self) -> dict:
        """Estadísticos de ajuste globales."""
        n = self.n_obs
        k = len(self.feature_names)
        ll = self.loglik

        # Log-likelihood nulo (sólo intercepto implícito)
        ll_null = -n * np.log(np.std(np.full(n, UPPER_LIMIT / 2)) + 1e-8)
        pseudo_r2 = 1 - ll / (ll_null + 1e-300)
        aic = -2 * ll + 2 * (k + 1)
        bic = -2 * ll + np.log(n) * (k + 1)

        # Wald test: H0: todos los coeficientes = 0
        try:
            beta = self.params[:-1]
            se = self.se[:-1]
            wald_stat = np.sum((beta / (se + 1e-300)) ** 2)
            wald_p = 1 - stats.chi2.cdf(wald_stat, df=k)
        except Exception:
            wald_stat, wald_p = np.nan, np.nan

        return {
            "N observaciones": n,
            "Log-Likelihood": round(ll, 4),
            "Pseudo R²": round(pseudo_r2, 4),
            "AIC": round(aic, 4),
            "BIC": round(bic, 4),
            "Sigma (σ)": round(self.sigma, 4),
            "Wald χ²": round(wald_stat, 4),
            "p-valor Wald": round(wald_p, 4),
        }

    def predict_latent(self, X: pd.DataFrame) -> np.ndarray:
        """Predicción de la variable latente E[y*|X]."""
        return X.values @ self.params[:-1]

    def predict_observed(self, X: pd.DataFrame) -> np.ndarray:
        """
        Predicción de la variable observada E[y|X].
        Fórmula: E[y|X] = σ[φ(a)-φ(b)] + Xβ[Φ(b)-Φ(a)] + lower*Φ(a) + upper*(1-Φ(b))
        donde a=(lower-Xβ)/σ, b=(upper-Xβ)/σ
        """
        xb = X.values @ self.params[:-1]
        sigma = self.sigma
        a = (self.lower - xb) / sigma
        b = (self.upper - xb) / sigma
        E_y = (
            sigma * (norm.pdf(a) - norm.pdf(b))
            + xb * (norm.cdf(b) - norm.cdf(a))
            + self.lower * norm.cdf(a)
            + self.upper * (1 - norm.cdf(b))
        )
        return E_y


# =====================================================
# 6. EFECTOS MARGINALES
# =====================================================

def compute_marginal_effects(
    model: DoubleCensoredTobit,
    X: pd.DataFrame,
) -> pd.DataFrame:
    """
    Calcula los efectos marginales del modelo Tobit de doble censura.

    Definiciones:
    - Efecto latente: β_j (coeficiente directo)
    - Efecto observado: β_j * [Φ(b) - Φ(a)]   (promedio muestral)
    - Efecto condicional: β_j * [1 - λ(a)*a - λ(b)*b]
      donde λ = φ/Φ (ratio de Mills inverso)
    """
    xb = model.predict_latent(X)
    sigma = model.sigma
    a = (model.lower - xb) / sigma
    b = (model.upper - xb) / sigma

    # Escalares promedio
    phi_a_mean = norm.pdf(a).mean()
    phi_b_mean = norm.pdf(b).mean()
    Phi_a_mean = norm.cdf(a).mean()
    Phi_b_mean = norm.cdf(b).mean()

    # Prob de no estar censurado
    prob_uncensored = (Phi_b_mean - Phi_a_mean)

    # Lambda (Mills) en media
    mill_a = phi_a_mean / (Phi_a_mean + 1e-300)
    mill_b = phi_b_mean / (1 - Phi_b_mean + 1e-300)

    beta = model.params[:-1]
    se = model.se[:-1] if model.se is not None else np.full_like(beta, np.nan)

    # Efectos
    eff_latent = beta
    eff_observed = beta * prob_uncensored
    eff_conditional = beta * (1 - mill_a * (a.mean()) - mill_b * (b.mean()))

    # Elasticidades en medias
    X_mean = X.mean()
    y_obs_mean = model.predict_observed(X).mean()
    elasticities = eff_observed * (X_mean.values / (y_obs_mean + 1e-300))

    df = pd.DataFrame(
        {
            "Variable": model.feature_names,
            "Efecto Latente (β)": eff_latent.round(4),
            "Efecto Observado": eff_observed.round(4),
            "Efecto Condicional": eff_conditional.round(4),
            "Elasticidad": elasticities.round(4),
            "Error Est. β": se.round(4),
        }
    )
    return df


# =====================================================
# 7. VISUALIZACIONES
# =====================================================

def plot_intensidad_hist(df: pd.DataFrame) -> plt.Figure:
    """Histograma de Intensidad mostrando acumulación en 0 y 100."""
    fig, ax = plt.subplots(figsize=(9, 5), dpi=150)
    y = df[REQUIRED_COL]
    interior = y[(y > LOWER_LIMIT) & (y < UPPER_LIMIT)]
    ax.hist(interior, bins=20, color=VERDE_CAFETO, alpha=0.7, edgecolor="white",
            label="Interior (0 < y < 100)")
    n_0 = (y <= LOWER_LIMIT).sum()
    n_100 = (y >= UPPER_LIMIT).sum()
    ax.axvline(LOWER_LIMIT, color=CAFE_OSCURO, linewidth=2.5, linestyle="--",
               label=f"Censurado en 0 (n={n_0})")
    ax.axvline(UPPER_LIMIT, color=DORADO, linewidth=2.5, linestyle="--",
               label=f"Censurado en 100 (n={n_100})")
    ax.set_xlabel("Intensidad de Adopción (%)", fontsize=12)
    ax.set_ylabel("Frecuencia", fontsize=12)
    ax.set_title("Distribución de Intensidad de Adopción PCS\n(Modelo Tobit — Doble Censura)",
                 fontsize=13, fontweight="bold", color=CAFE_OSCURO)
    ax.legend(framealpha=0.9)
    ax.grid(axis="y", alpha=0.3)
    fig.tight_layout()
    return fig


def plot_boxplot_municipio(df: pd.DataFrame) -> plt.Figure:
    """Boxplot de Intensidad por municipio."""
    fig, ax = plt.subplots(figsize=(10, 5), dpi=150)
    if "Mpio" not in df.columns:
        ax.text(0.5, 0.5, "Variable 'Mpio' no disponible", ha="center", va="center",
                transform=ax.transAxes, fontsize=13)
        return fig
    palette = [VERDE_CAFETO, CAFE_OSCURO, DORADO, "#66BB6A", "#8D6E63"]
    mpios = df["Mpio"].unique()
    data_by_mpio = [df[df["Mpio"] == m][REQUIRED_COL].dropna() for m in mpios]
    bp = ax.boxplot(data_by_mpio, labels=mpios, patch_artist=True,
                    medianprops=dict(color="white", linewidth=2))
    for patch, color in zip(bp["boxes"], palette * 10):
        patch.set_facecolor(color)
        patch.set_alpha(0.8)
    ax.set_xlabel("Municipio", fontsize=12)
    ax.set_ylabel("Intensidad (%)", fontsize=12)
    ax.set_title("Intensidad de Adopción PCS por Municipio — Huila",
                 fontsize=13, fontweight="bold", color=CAFE_OSCURO)
    ax.grid(axis="y", alpha=0.3)
    fig.tight_layout()
    return fig


def plot_violin(df: pd.DataFrame) -> plt.Figure:
    """Violin plot de Intensidad por municipio."""
    fig, ax = plt.subplots(figsize=(10, 5), dpi=150)
    if "Mpio" not in df.columns:
        ax.text(0.5, 0.5, "Variable 'Mpio' no disponible", ha="center", va="center",
                transform=ax.transAxes)
        return fig
    mpios = sorted(df["Mpio"].unique())
    palette = sns.color_palette([VERDE_CAFETO, CAFE_OSCURO, DORADO, "#66BB6A", "#8D6E63"])
    data_plot = [(df[df["Mpio"] == m][REQUIRED_COL].dropna().values) for m in mpios]
    parts = ax.violinplot(data_plot, positions=range(len(mpios)),
                          showmedians=True, showextrema=True)
    for i, pc in enumerate(parts["bodies"]):
        pc.set_facecolor(palette[i % len(palette)])
        pc.set_alpha(0.75)
    ax.set_xticks(range(len(mpios)))
    ax.set_xticklabels(mpios, fontsize=10)
    ax.set_ylabel("Intensidad (%)", fontsize=12)
    ax.set_title("Violin Plot — Intensidad de Adopción por Municipio",
                 fontsize=13, fontweight="bold", color=CAFE_OSCURO)
    ax.grid(axis="y", alpha=0.3)
    fig.tight_layout()
    return fig


def plot_heatmap(df_model: pd.DataFrame, feature_cols: list[str]) -> plt.Figure:
    """Heatmap de correlaciones entre predictores y variable dependiente."""
    cols = [REQUIRED_COL] + [c for c in feature_cols if c in df_model.columns][:12]
    corr = df_model[cols].corr()
    fig, ax = plt.subplots(figsize=(min(len(cols), 12) + 1, min(len(cols), 12)), dpi=150)
    mask = np.zeros_like(corr, dtype=bool)
    np.fill_diagonal(mask, True)
    cmap = sns.diverging_palette(145, 20, as_cmap=True)
    sns.heatmap(corr, ax=ax, mask=mask, cmap=cmap, center=0,
                annot=True, fmt=".2f", annot_kws={"size": 8},
                linewidths=0.5, square=True, cbar_kws={"shrink": 0.8})
    ax.set_title("Matriz de Correlaciones — Variables del Modelo",
                 fontsize=12, fontweight="bold", color=CAFE_OSCURO, pad=12)
    fig.tight_layout()
    return fig


def plot_forest(marginal_df: pd.DataFrame) -> plt.Figure:
    """Forest plot para efectos marginales observados."""
    df = marginal_df.copy().sort_values("Efecto Observado")
    fig, ax = plt.subplots(figsize=(9, max(5, len(df) * 0.5 + 1)), dpi=150)

    colors = [VERDE_CAFETO if v >= 0 else CAFE_OSCURO for v in df["Efecto Observado"]]
    ci = df["Error Est. β"] * 1.96 * abs(
        df["Efecto Observado"] / (df["Efecto Latente (β)"].replace(0, np.nan))
    ).fillna(1)

    ax.barh(df["Variable"], df["Efecto Observado"], color=colors, alpha=0.75,
            xerr=ci, capsize=3, error_kw={"elinewidth": 1, "ecolor": "gray"})
    ax.axvline(0, color="black", linewidth=0.8, linestyle="--")
    ax.set_xlabel("Efecto Marginal Observado", fontsize=11)
    ax.set_title("Forest Plot — Efectos Marginales Observados (Modelo Tobit)",
                 fontsize=12, fontweight="bold", color=CAFE_OSCURO)
    ax.grid(axis="x", alpha=0.3)
    pos_patch = mpatches.Patch(color=VERDE_CAFETO, label="Efecto positivo")
    neg_patch = mpatches.Patch(color=CAFE_OSCURO, label="Efecto negativo")
    ax.legend(handles=[pos_patch, neg_patch], fontsize=9)
    fig.tight_layout()
    return fig


def plot_obs_vs_pred(y_true: np.ndarray, y_pred: np.ndarray) -> plt.Figure:
    """Gráfico observado vs predicho."""
    fig, ax = plt.subplots(figsize=(7, 6), dpi=150)
    ax.scatter(y_pred, y_true, color=VERDE_CAFETO, alpha=0.5, s=25, edgecolors="none")
    lim_min = min(y_true.min(), y_pred.min()) - 2
    lim_max = max(y_true.max(), y_pred.max()) + 2
    ax.plot([lim_min, lim_max], [lim_min, lim_max], "r--", linewidth=1.2, label="45°")
    ax.set_xlabel("Valores Predichos E[y|X]", fontsize=11)
    ax.set_ylabel("Valores Observados", fontsize=11)
    ax.set_title("Predicho vs Observado — Modelo Tobit",
                 fontsize=12, fontweight="bold", color=CAFE_OSCURO)
    ax.legend(fontsize=9)
    ax.grid(alpha=0.3)
    rmse = np.sqrt(mean_squared_error(y_true, y_pred))
    ax.annotate(f"RMSE={rmse:.3f}", xy=(0.05, 0.92), xycoords="axes fraction",
                fontsize=10, color=CAFE_OSCURO)
    fig.tight_layout()
    return fig


def plot_residuals(y_true: np.ndarray, y_pred: np.ndarray) -> plt.Figure:
    """Gráfico de residuos."""
    residuals = y_true - y_pred
    fig, axes = plt.subplots(1, 2, figsize=(12, 5), dpi=150)

    axes[0].scatter(y_pred, residuals, color=VERDE_CAFETO, alpha=0.5, s=25)
    axes[0].axhline(0, color="red", linewidth=1, linestyle="--")
    axes[0].set_xlabel("Valores Predichos", fontsize=11)
    axes[0].set_ylabel("Residuos", fontsize=11)
    axes[0].set_title("Residuos vs Predichos", fontsize=12, fontweight="bold",
                      color=CAFE_OSCURO)
    axes[0].grid(alpha=0.3)

    stats.probplot(residuals, dist="norm", plot=axes[1])
    axes[1].get_lines()[0].set(color=VERDE_CAFETO, markersize=4, alpha=0.6)
    axes[1].get_lines()[1].set(color="red", linewidth=1.5)
    axes[1].set_title("QQ Plot — Normalidad de Residuos", fontsize=12,
                      fontweight="bold", color=CAFE_OSCURO)
    axes[1].grid(alpha=0.3)

    fig.tight_layout()
    return fig


def plot_ols_vs_tobit(
    X: pd.DataFrame,
    y: pd.Series,
    tobit: DoubleCensoredTobit,
) -> plt.Figure:
    """Comparación gráfica OLS vs Tobit para primer predictor."""
    fig, ax = plt.subplots(figsize=(8, 5), dpi=150)

    # Usar primer predictor (excluir constante si hay)
    feat_col = X.columns[0]
    x_range = np.linspace(X[feat_col].min(), X[feat_col].max(), 200)
    X_mean = X.mean()

    # Predicciones Tobit variando primer predictor
    X_sim = pd.DataFrame(
        np.tile(X_mean.values, (200, 1)), columns=X.columns
    )
    X_sim[feat_col] = x_range
    tobit_pred = tobit.predict_observed(X_sim)

    # OLS
    ols_result = sm.OLS(y, X).fit()
    X_sim_np = X_sim.values
    ols_pred = X_sim_np @ ols_result.params

    ax.scatter(X[feat_col], y, color=VERDE_CAFETO, alpha=0.35, s=20, label="Observados")
    ax.plot(x_range, tobit_pred, color=DORADO, linewidth=2.5, label="Tobit E[y|X]")
    ax.plot(x_range, ols_pred, color=CAFE_OSCURO, linewidth=2, linestyle="--",
            label="OLS")
    ax.axhline(LOWER_LIMIT, color="gray", linewidth=0.8, linestyle=":")
    ax.axhline(UPPER_LIMIT, color="gray", linewidth=0.8, linestyle=":")
    ax.set_xlabel(feat_col, fontsize=11)
    ax.set_ylabel("Intensidad (%)", fontsize=11)
    ax.set_title(f"OLS vs Tobit — {feat_col}",
                 fontsize=12, fontweight="bold", color=CAFE_OSCURO)
    ax.legend(fontsize=9)
    ax.grid(alpha=0.3)
    fig.tight_layout()
    return fig


# =====================================================
# 8. EXPORTACIÓN WORD
# =====================================================

def set_cell_background(cell, hex_color: str) -> None:
    """Establece color de fondo en celda Word."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color.lstrip("#"))
    tcPr.append(shd)


def generate_word_report(
    df: pd.DataFrame,
    model: DoubleCensoredTobit,
    summary: pd.DataFrame,
    marginal: pd.DataFrame,
    fit_stats: dict,
    feature_cols: list[str],
) -> bytes:
    """
    Genera reporte Word en formato APA 7.
    Retorna bytes del documento .docx.
    """
    doc = Document()

    # Estilos
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # Márgenes APA: 1 pulgada
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # ---- PORTADA ----
    doc.add_paragraph()
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_p.add_run(
        "Determinantes de la Intensidad de Adopción de Prácticas\n"
        "de Conservación de Suelos en Fincas Cafeteras del Huila:\n"
        "Un Análisis mediante Modelo Tobit de Doble Censura"
    )
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x2E, 0x7D, 0x32)

    doc.add_paragraph()
    inst_p = doc.add_paragraph()
    inst_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    inst_p.add_run(
        "Centro Nacional de Investigaciones de Café (Cenicafé)\n"
        "Departamento del Huila, Colombia\n"
        "2026"
    ).font.size = Pt(12)
    doc.add_page_break()

    # ---- RESUMEN ----
    doc.add_heading("Resumen Ejecutivo", level=1)
    n_obs = fit_stats.get("N observaciones", len(df))
    doc.add_paragraph(
        f"El presente estudio analiza los determinantes socioeconómicos, productivos y "
        f"tecnológicos de la intensidad de adopción de prácticas de conservación de suelos "
        f"(PCS) en {n_obs} fincas cafeteras del departamento del Huila (Colombia), "
        f"utilizando un modelo Tobit de doble censura con límites en 0 y 100. "
        f"La variable dependiente (Intensidad) representa el porcentaje de adopción y "
        f"presenta acumulación en los extremos, justificando el enfoque Tobit. "
        f"Los resultados permiten identificar los factores que determinan tanto la "
        f"probabilidad de adopción como la intensidad del proceso."
    )
    doc.add_paragraph(
        f"Palabras clave: Tobit, censura, prácticas de conservación de suelos, "
        f"adopción tecnológica, Huila, café."
    )

    # ---- METODOLOGÍA ----
    doc.add_heading("Metodología", level=1)
    doc.add_heading("Especificación del Modelo Tobit", level=2)
    doc.add_paragraph(
        "El modelo Tobit de doble censura (Tobin, 1958; Amemiya, 1984) especifica "
        "la siguiente estructura latente:"
    )
    doc.add_paragraph("y*ᵢ = Xᵢβ + εᵢ,    εᵢ ~ N(0, σ²)", style="Normal")
    doc.add_paragraph(
        "La variable observada se relaciona con la latente de la siguiente forma:\n"
        "  yᵢ = 0        si y*ᵢ ≤ 0\n"
        "  yᵢ = y*ᵢ      si 0 < y*ᵢ < 100\n"
        "  yᵢ = 100      si y*ᵢ ≥ 100\n\n"
        "IMPORTANTE: Los coeficientes β estiman el efecto sobre la variable LATENTE, "
        "no sobre la variable observada. Los efectos marginales observados requieren "
        "ajuste por la probabilidad de no censura."
    )
    doc.add_paragraph(
        f"La estimación se realizó por Máxima Verosimilitud (MLE). "
        f"Se utilizaron {n_obs} observaciones de fincas cafeteras en cinco municipios "
        f"del Huila: Gigante, Garzón, Acevedo, San Agustín y La Plata."
    )

    # ---- RESULTADOS ----
    doc.add_heading("Resultados", level=1)
    doc.add_heading("Estadísticos de Ajuste", level=2)

    # Tabla fit stats
    t = doc.add_table(rows=1, cols=2)
    t.style = "Table Grid"
    hdr = t.rows[0].cells
    hdr[0].text = "Estadístico"
    hdr[1].text = "Valor"
    for cell in hdr:
        set_cell_background(cell, "2E7D32")
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    for k, v in fit_stats.items():
        row = t.add_row().cells
        row[0].text = str(k)
        row[1].text = str(v)

    doc.add_paragraph()
    doc.add_heading("Coeficientes del Modelo Tobit", level=2)
    doc.add_paragraph(
        "Nota: *** p<0.01; ** p<0.05; * p<0.10. Los coeficientes representan "
        "efectos sobre la variable latente E[y*|X]."
    )

    # Tabla coeficientes
    cols_show = ["Variable", "Coeficiente", "Error Estándar", "z-Valor",
                 "p-Valor", "Sig.", "IC 95% Inf.", "IC 95% Sup."]
    t2 = doc.add_table(rows=1, cols=len(cols_show))
    t2.style = "Table Grid"
    hdr2 = t2.rows[0].cells
    for i, col in enumerate(cols_show):
        hdr2[i].text = col
        set_cell_background(hdr2[i], "2E7D32")
        for p in hdr2[i].paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    for _, row_data in summary.iterrows():
        row = t2.add_row().cells
        for i, col in enumerate(cols_show):
            row[i].text = str(row_data[col])

    # Efectos marginales
    doc.add_paragraph()
    doc.add_heading("Efectos Marginales", level=2)
    doc.add_paragraph(
        "Los efectos marginales observados miden el cambio esperado en la variable "
        "observada (Intensidad) ante un cambio unitario en el predictor, ponderado "
        "por la probabilidad de pertenecer al rango no censurado."
    )
    cols_m = list(marginal.columns)
    t3 = doc.add_table(rows=1, cols=len(cols_m))
    t3.style = "Table Grid"
    hdr3 = t3.rows[0].cells
    for i, col in enumerate(cols_m):
        hdr3[i].text = col
        set_cell_background(hdr3[i], "4E342E")
        for p in hdr3[i].paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    for _, row_data in marginal.iterrows():
        row = t3.add_row().cells
        for i, col in enumerate(cols_m):
            row[i].text = str(row_data[col])

    # ---- CONCLUSIONES ----
    doc.add_page_break()
    doc.add_heading("Conclusiones", level=1)
    sig_vars = summary[summary["p-Valor"].astype(float) < 0.05]["Variable"].tolist()
    doc.add_paragraph(
        f"El modelo Tobit de doble censura ajustado presenta un pseudo-R² de "
        f"{fit_stats.get('Pseudo R²', 'N/D')}, con un AIC de {fit_stats.get('AIC', 'N/D')} "
        f"y BIC de {fit_stats.get('BIC', 'N/D')}. "
        f"Las variables estadísticamente significativas (p<0.05) son: "
        f"{', '.join(sig_vars) if sig_vars else 'ninguna al nivel del 5%'}. "
        f"El valor de sigma (σ={fit_stats.get('Sigma (σ)', 'N/D')}) refleja la varianza "
        f"del término de error del modelo latente."
    )

    # ---- REFERENCIAS ----
    doc.add_heading("Referencias", level=1)
    refs = [
        "Amemiya, T. (1984). Tobit models: A survey. Journal of Econometrics, 24(1–2), 3–61. "
        "https://doi.org/10.1016/0304-4076(84)90074-5",
        "Tobin, J. (1958). Estimation of relationships for limited dependent variables. "
        "Econometrica, 26(1), 24–36. https://doi.org/10.2307/1907382",
        "Greene, W. H. (2018). Econometric Analysis (8th ed.). Pearson.",
        "Wooldridge, J. M. (2010). Econometric Analysis of Cross Section and Panel Data. MIT Press.",
        "Cenicafé. (2026). Adopción de Prácticas de Conservación de Suelos en el Huila. "
        "Centro Nacional de Investigaciones de Café.",
    ]
    for ref in refs:
        p = doc.add_paragraph(ref, style="Normal")
        p.paragraph_format.left_indent = Inches(0.5)
        p.paragraph_format.first_line_indent = Inches(-0.5)

    # Serializar
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()


# =====================================================
# 9. EXPORTACIÓN EXCEL
# =====================================================

def generate_excel_report(
    df: pd.DataFrame,
    df_model: pd.DataFrame,
    summary: pd.DataFrame,
    marginal: pd.DataFrame,
    fit_stats: dict,
    y_pred: np.ndarray,
    y_true: np.ndarray,
    feature_cols: list[str],
) -> bytes:
    """Genera reporte Excel multi-hoja con formato profesional."""
    buf = io.BytesIO()
    with pd.ExcelWriter(buf, engine="openpyxl") as writer:
        # Colores
        header_fill = PatternFill("solid", fgColor="2E7D32")
        alt_fill = PatternFill("solid", fgColor="F1F8E9")
        header_font = Font(bold=True, color="FFFFFF", name="Calibri", size=11)
        body_font = Font(name="Calibri", size=10)
        center = Alignment(horizontal="center", vertical="center", wrap_text=True)
        left = Alignment(horizontal="left", vertical="center")
        thin = Side(style="thin", color="CCCCCC")
        border = Border(left=thin, right=thin, top=thin, bottom=thin)

        def style_ws(ws, df_sheet: pd.DataFrame) -> None:
            """Aplica formato al worksheet."""
            for col_idx, col_name in enumerate(df_sheet.columns, 1):
                cell = ws.cell(row=1, column=col_idx, value=col_name)
                cell.fill = header_fill
                cell.font = header_font
                cell.alignment = center
                cell.border = border
            for row_idx, row in enumerate(df_sheet.itertuples(index=False), 2):
                fill = alt_fill if row_idx % 2 == 0 else PatternFill()
                for col_idx, value in enumerate(row, 1):
                    cell = ws.cell(row=row_idx, column=col_idx, value=value)
                    cell.font = body_font
                    cell.alignment = left
                    cell.border = border
                    cell.fill = fill
            for col in ws.columns:
                max_len = max(len(str(c.value or "")) for c in col)
                ws.column_dimensions[get_column_letter(col[0].column)].width = min(max_len + 4, 30)

        # Hoja 1 — Datos_Modelo
        df_model.to_excel(writer, sheet_name="Datos_Modelo", index=False)
        style_ws(writer.sheets["Datos_Modelo"], df_model)

        # Hoja 2 — Resultados_Tobit
        summary.to_excel(writer, sheet_name="Resultados_Tobit", index=False)
        style_ws(writer.sheets["Resultados_Tobit"], summary)

        # Hoja 3 — Efectos_Marginales
        marginal.to_excel(writer, sheet_name="Efectos_Marginales", index=False)
        style_ws(writer.sheets["Efectos_Marginales"], marginal)

        # Hoja 4 — Estadisticas
        avail_cols = [c for c in feature_cols + [REQUIRED_COL] if c in df_model.columns]
        desc = df_model[avail_cols].describe().T.reset_index()
        desc.columns = ["Variable"] + list(desc.columns[1:])
        fit_df = pd.DataFrame(list(fit_stats.items()), columns=["Estadístico", "Valor"])
        fit_df.to_excel(writer, sheet_name="Estadisticas", index=False, startrow=0)
        desc.to_excel(writer, sheet_name="Estadisticas", index=False,
                      startrow=len(fit_df) + 3)
        style_ws(writer.sheets["Estadisticas"], fit_df)

        # Hoja 5 — Predicciones
        pred_df = pd.DataFrame({
            "Observado": y_true,
            "Predicho": y_pred,
            "Residuo": y_true - y_pred,
        })
        pred_df.to_excel(writer, sheet_name="Predicciones", index=False)
        style_ws(writer.sheets["Predicciones"], pred_df)

        # Hoja 6 — Residuos
        residuos_df = pd.DataFrame({
            "Residuo": y_true - y_pred,
            "Residuo_Abs": np.abs(y_true - y_pred),
            "Residuo_Sq": (y_true - y_pred) ** 2,
        })
        residuos_df.to_excel(writer, sheet_name="Residuos", index=False)
        style_ws(writer.sheets["Residuos"], residuos_df)

    buf.seek(0)
    return buf.getvalue()


# =====================================================
# 10. SIDEBAR
# =====================================================

def render_sidebar() -> dict:
    """Renderiza la barra lateral y retorna opciones seleccionadas."""
    with st.sidebar:
        st.image(
            "https://via.placeholder.com/220x60/2E7D32/FFF8E1?text=☕+Cenicafé",
            use_column_width=True,
        )
        st.markdown("---")
        st.markdown(
            "<h3 style='color:#F9A825;font-size:16px;'>🔬 Análisis Tobit PCS</h3>",
            unsafe_allow_html=True,
        )
        st.markdown(
            "<p style='font-size:12px;color:#FFF8E1;'>"
            "Modelo Tobit de Doble Censura<br>"
            "Intensidad de Adopción PCS<br>"
            "Huila · 422 fincas cafeteras"
            "</p>",
            unsafe_allow_html=True,
        )
        st.markdown("---")

        st.markdown("<b style='color:#F9A825;'>⚙️ Opciones del Modelo</b>",
                    unsafe_allow_html=True)
        winsorize = st.checkbox("Winsorización de variables (5%)", value=False)
        add_constant = st.checkbox("Agregar constante al modelo", value=True)
        alpha = st.selectbox("Nivel de significancia", [0.01, 0.05, 0.10], index=1)

        st.markdown("---")
        st.markdown("<b style='color:#F9A825;'>📊 Visualización</b>",
                    unsafe_allow_html=True)
        max_vars_heatmap = st.slider("Variables en heatmap", 5, 20, 12)

        st.markdown("---")
        st.markdown(
            "<p style='font-size:10px;color:#FFF8E1;text-align:center;'>"
            "v1.0 · Cenicafé 2026<br>"
            "Python 3.11 · Streamlit"
            "</p>",
            unsafe_allow_html=True,
        )

    return {
        "winsorize": winsorize,
        "add_constant": add_constant,
        "alpha": alpha,
        "max_vars_heatmap": max_vars_heatmap,
    }


# =====================================================
# 11. PESTAÑAS STREAMLIT
# =====================================================

def main() -> None:
    """Función principal de la aplicación."""

    # Header
    col_logo, col_title = st.columns([1, 5])
    with col_logo:
        st.markdown(
            "<div style='font-size:52px;text-align:center;'>☕</div>",
            unsafe_allow_html=True,
        )
    with col_title:
        st.markdown(
            f"<h1 style='color:{VERDE_CAFETO};margin-bottom:0;'>"
            "Modelo Tobit de Doble Censura</h1>"
            f"<p style='color:{CAFE_OSCURO};font-size:15px;margin-top:4px;'>"
            "Intensidad de Adopción de Prácticas de Conservación de Suelos · "
            "Cenicafé · Huila, Colombia</p>",
            unsafe_allow_html=True,
        )
    divider()

    # Sidebar
    opts = render_sidebar()

    # Tabs
    tabs = st.tabs([
        "📊 Datos Generales",
        "🔬 Modelado Tobit",
        "📂 Carga de Datos",
        "📈 Visualizaciones",
        "📄 Exportación",
        "🚀 GitHub & Deploy",
    ])

    # ============================================================
    # Estado de sesión para datos
    # ============================================================
    if "df" not in st.session_state:
        with st.spinner("Cargando dataset..."):
            df_default = load_default_data()
            if df_default is not None:
                st.session_state.df = df_default
                st.session_state.data_source = DEFAULT_FILE

    df: Optional[pd.DataFrame] = st.session_state.get("df", None)

    # ============================================================
    # TAB 1 — DATOS GENERALES
    # ============================================================
    with tabs[0]:
        st.header("📊 Datos Generales y Estadísticas Descriptivas")

        if df is None:
            st.info("⬆️ Cargue un dataset en la pestaña **Carga de Datos** para comenzar.")
            st.stop()

        valid, errors = validate_dataset(df)
        if not valid:
            for e in errors:
                st.error(e)
            st.stop()

        st.success(
            f"✅ Dataset validado: **{len(df):,} observaciones** · "
            f"**{df.shape[1]} columnas** · "
            f"Fuente: {st.session_state.get('data_source', 'Cargado por usuario')}"
        )
        divider()

        # KPIs
        y = df[REQUIRED_COL]
        metric_row([
            {"label": "N Observaciones", "value": f"{len(df):,}"},
            {"label": "Intensidad Media (%)", "value": f"{y.mean():.2f}"},
            {"label": "Desv. Estándar", "value": f"{y.std():.2f}"},
            {"label": "Censurados en 0", "value": f"{(y <= 0).sum()}"},
            {"label": "Censurados en 100", "value": f"{(y >= 100).sum()}"},
            {"label": "No censurados", "value": f"{((y > 0) & (y < 100)).sum()}"},
        ])
        divider()

        # Municipios
        if "Mpio" in df.columns:
            st.subheader("Frecuencias por Municipio")
            freq = df["Mpio"].value_counts().reset_index()
            freq.columns = ["Municipio", "N"]
            freq["% del Total"] = (freq["N"] / len(df) * 100).round(2)
            col1, col2 = st.columns([1, 2])
            with col1:
                st.dataframe(styled_df(freq), use_container_width=True)
            with col2:
                fig_bar = px.bar(
                    freq, x="Municipio", y="N",
                    color="N", color_continuous_scale=[[0, CREMA], [1, VERDE_CAFETO]],
                    text="N", title="Distribución de Fincas por Municipio",
                )
                fig_bar.update_traces(textposition="outside")
                fig_bar.update_layout(
                    showlegend=False, plot_bgcolor="white",
                    title_font_color=CAFE_OSCURO,
                )
                st.plotly_chart(fig_bar, use_container_width=True)

        divider()

        # Distribución Intensidad
        st.subheader("Distribución de la Variable Dependiente: Intensidad")
        col_h, col_v = st.columns(2)
        with col_h:
            fig_hist = plot_intensidad_hist(df)
            st.pyplot(fig_hist, use_container_width=True)
            plt.close(fig_hist)
        with col_v:
            if "Mpio" in df.columns:
                fig_vio = plot_violin(df)
                st.pyplot(fig_vio, use_container_width=True)
                plt.close(fig_vio)

        # Boxplot
        if "Mpio" in df.columns:
            fig_box = plot_boxplot_municipio(df)
            st.pyplot(fig_box, use_container_width=True)
            plt.close(fig_box)

        divider()

        # Estadísticas descriptivas
        st.subheader("Estadísticas Descriptivas")
        num_cols = df.select_dtypes(include=np.number).columns.tolist()
        desc = df[num_cols].describe().T.round(4)
        st.dataframe(desc, use_container_width=True)

        divider()

        # Heatmap
        st.subheader("Matriz de Correlaciones")
        try:
            _, y_ser, feat_cols = preprocess(df, winsorize=opts["winsorize"])
            n_vars = min(opts["max_vars_heatmap"], len(feat_cols))
            top_corr_cols = (
                df[feat_cols].corrwith(df[REQUIRED_COL])
                .abs().nlargest(n_vars).index.tolist()
            )
            df_heat = df[top_corr_cols + [REQUIRED_COL]].dropna()
            fig_hm = plot_heatmap(df_heat, top_corr_cols)
            st.pyplot(fig_hm, use_container_width=True)
            plt.close(fig_hm)
        except Exception as exc:
            st.warning(f"No se pudo generar el heatmap: {exc}")

        divider()

        # Tabla interactiva
        st.subheader("Tabla Interactiva del Dataset")
        st.dataframe(df.head(50), use_container_width=True, height=400)

    # ============================================================
    # TAB 2 — MODELADO TOBIT
    # ============================================================
    with tabs[1]:
        st.header("🔬 Modelado Tobit de Doble Censura")

        if df is None:
            st.info("⬆️ Cargue un dataset primero.")
            st.stop()

        info_box(
            "Los coeficientes del modelo Tobit estiman el efecto sobre la "
            "<b>variable latente</b> E[y*|X], no sobre la variable observada. "
            "Los efectos marginales observados requieren corrección por la "
            "probabilidad de no censura."
        )

        with st.expander("⚙️ Paso 1 — Preparación de datos", expanded=True):
            try:
                df_model, y_ser, feat_cols = preprocess(df, winsorize=opts["winsorize"])
                X = df_model[feat_cols].copy()
                if opts["add_constant"] and "const" not in X.columns:
                    X.insert(0, "const", 1.0)
                st.success(
                    f"✅ Variables predictoras: **{len(feat_cols)}** · "
                    f"Observaciones: **{len(y_ser):,}** · "
                    f"Constante: {'Sí' if opts['add_constant'] else 'No'}"
                )
                st.write("**Primeras filas de X:**")
                st.dataframe(X.head(5), use_container_width=True)
            except Exception as exc:
                st.error(f"Error en preprocesamiento: {exc}")
                st.stop()

        with st.expander("🚀 Paso 2 — Estimación del Modelo", expanded=True):
            if st.button("▶️ Ejecutar Modelo Tobit", type="primary"):
                with st.spinner("Estimando por Máxima Verosimilitud..."):
                    try:
                        tobit = DoubleCensoredTobit(lower=LOWER_LIMIT, upper=UPPER_LIMIT)
                        tobit.fit(X, y_ser)
                        st.session_state.tobit = tobit
                        st.session_state.X = X
                        st.session_state.y_ser = y_ser
                        st.session_state.feat_cols = feat_cols
                        st.success("✅ Modelo estimado exitosamente.")
                    except Exception as exc:
                        st.error(f"Error en estimación: {exc}")
                        st.stop()
            elif "tobit" not in st.session_state:
                st.info("Presione el botón para estimar el modelo.")

        tobit: Optional[DoubleCensoredTobit] = st.session_state.get("tobit")
        if tobit is None:
            st.stop()

        X_model = st.session_state.X
        y_model = st.session_state.y_ser

        with st.expander("📋 Paso 3 — Estadísticos de Ajuste", expanded=True):
            fit_stats = tobit.fit_stats()
            st.session_state.fit_stats = fit_stats
            col1, col2, col3, col4 = st.columns(4)
            col1.metric("Log-Likelihood", f"{fit_stats['Log-Likelihood']:.4f}")
            col2.metric("Pseudo R²", f"{fit_stats['Pseudo R²']:.4f}")
            col3.metric("AIC", f"{fit_stats['AIC']:.4f}")
            col4.metric("BIC", f"{fit_stats['BIC']:.4f}")
            col5, col6, col7, col8 = st.columns(4)
            col5.metric("Sigma (σ)", f"{fit_stats['Sigma (σ)']:.4f}")
            col6.metric("Wald χ²", f"{fit_stats['Wald χ²']:.4f}")
            col7.metric("p-valor Wald", f"{fit_stats['p-valor Wald']:.4f}")
            col8.metric("N", f"{fit_stats['N observaciones']:,}")

        with st.expander("📊 Paso 4 — Coeficientes del Modelo", expanded=True):
            try:
                summary_df = tobit.summary_df()
                st.session_state.summary_df = summary_df
                alpha_val = opts["alpha"]
                sig_mask = summary_df["p-Valor"].astype(float) < alpha_val
                st.dataframe(
                    summary_df.style.apply(
                        lambda x: [
                            f"background-color:{CREMA};font-weight:bold"
                            if sig_mask.iloc[i] else ""
                            for i in range(len(x))
                        ],
                        axis=0,
                    ),
                    use_container_width=True,
                    height=400,
                )
                n_sig = sig_mask.sum()
                st.info(
                    f"Variables significativas (p < {alpha_val}): **{n_sig}** de {len(summary_df)}"
                )
            except Exception as exc:
                st.error(f"Error en tabla de coeficientes: {exc}")

        with st.expander("📐 Paso 5 — Efectos Marginales", expanded=True):
            try:
                marginal_df = compute_marginal_effects(tobit, X_model)
                st.session_state.marginal_df = marginal_df
                st.dataframe(styled_df(marginal_df), use_container_width=True)
                st.caption(
                    "Efecto Latente = β; Efecto Observado = β × P(no censurado); "
                    "Efecto Condicional = β × [1 − λ(a)·a − λ(b)·b]; "
                    "Elasticidad = Efecto Observado × (x̄ / ȳ)"
                )
            except Exception as exc:
                st.error(f"Error en efectos marginales: {exc}")

        with st.expander("🧠 Paso 6 — Interpretación Automática", expanded=False):
            try:
                summary_df = st.session_state.get("summary_df", tobit.summary_df())
                marginal_df = st.session_state.get("marginal_df",
                                                    compute_marginal_effects(tobit, X_model))
                sig = summary_df[summary_df["p-Valor"].astype(float) < opts["alpha"]]
                pos_vars = sig[sig["Coeficiente"].astype(float) > 0]["Variable"].tolist()
                neg_vars = sig[sig["Coeficiente"].astype(float) < 0]["Variable"].tolist()
                st.markdown(f"""
**Interpretación automática — Nivel α = {opts['alpha']}**

- **Variables con efecto positivo y significativo:** {', '.join(pos_vars) if pos_vars else 'Ninguna'}
- **Variables con efecto negativo y significativo:** {', '.join(neg_vars) if neg_vars else 'Ninguna'}
- **Sigma (σ = {tobit.sigma:.4f}):** Indica la variabilidad del término de error en el modelo latente.
- **Pseudo R² = {fit_stats['Pseudo R²']:.4f}:** Medida de bondad de ajuste basada en log-verosimilitud.
- **Test de Wald (χ² = {fit_stats['Wald χ²']:.4f}, p = {fit_stats['p-valor Wald']:.4f}):**
  {'El modelo es globalmente significativo.' if fit_stats['p-valor Wald'] < 0.05 else 'El modelo no es globalmente significativo al 5%.'}

> ⚠️ **Recordatorio metodológico:** Los coeficientes β del modelo Tobit estiman efectos sobre
> la variable latente y*, no sobre la Intensidad observada. Para interpretar impactos reales,
> utilice los efectos marginales observados de la tabla anterior.
                """)
            except Exception as exc:
                st.warning(f"No se pudo generar interpretación: {exc}")

    # ============================================================
    # TAB 3 — CARGA DE DATOS
    # ============================================================
    with tabs[2]:
        st.header("📂 Carga de Datos")
        st.markdown(
            "Suba su archivo CSV o Excel. Si existe "
            f"`{DEFAULT_FILE}` en el directorio, se carga automáticamente."
        )

        # Estado actual
        if df is not None:
            st.success(
                f"✅ Dataset activo: {st.session_state.get('data_source', '—')} "
                f"({len(df):,} filas × {df.shape[1]} cols)"
            )

        divider()

        sep_options = {";": ";", ",": ",", "Tab (\\t)": "\t", "|": "|"}
        sep_label = st.selectbox("Separador del CSV", list(sep_options.keys()))
        sep = sep_options[sep_label]

        uploaded = st.file_uploader(
            "Seleccione archivo CSV o Excel (.xlsx)",
            type=["csv", "xlsx"],
        )

        if uploaded is not None:
            try:
                if uploaded.name.endswith(".xlsx"):
                    df_new = pd.read_excel(uploaded)
                else:
                    df_new = pd.read_csv(uploaded, sep=sep)

                st.info(f"Archivo: **{uploaded.name}** · {len(df_new):,} filas · {df_new.shape[1]} cols")
                st.subheader("Vista previa")
                st.dataframe(df_new.head(10), use_container_width=True)

                st.subheader("Tipos de datos")
                types_df = pd.DataFrame({
                    "Columna": df_new.dtypes.index,
                    "Tipo": df_new.dtypes.values,
                    "Nulos": df_new.isnull().sum().values,
                    "% Nulos": (df_new.isnull().mean() * 100).round(2).values,
                })
                st.dataframe(types_df, use_container_width=True)

                # Validar
                valid_new, err_new = validate_dataset(df_new)
                if valid_new:
                    if st.button("✅ Usar este dataset", type="primary"):
                        st.session_state.df = df_new
                        st.session_state.data_source = uploaded.name
                        # Limpiar modelo anterior
                        for key in ["tobit", "X", "y_ser", "feat_cols",
                                    "fit_stats", "summary_df", "marginal_df"]:
                            st.session_state.pop(key, None)
                        st.success("Dataset cargado. Vaya a la pestaña Modelado Tobit.")
                        st.rerun()
                else:
                    for e in err_new:
                        st.error(e)
            except Exception as exc:
                st.error(f"Error al leer el archivo: {exc}")

    # ============================================================
    # TAB 4 — VISUALIZACIONES CIENTÍFICAS
    # ============================================================
    with tabs[3]:
        st.header("📈 Visualizaciones Científicas")

        if df is None:
            st.info("⬆️ Cargue un dataset primero.")
            st.stop()

        tobit = st.session_state.get("tobit")
        X_model = st.session_state.get("X")
        y_model = st.session_state.get("y_ser")

        # Histograma censurado
        st.subheader("1. Histograma — Distribución Censurada de Intensidad")
        fig1 = plot_intensidad_hist(df)
        st.pyplot(fig1, use_container_width=True)
        plt.close(fig1)

        divider()

        if tobit is not None and X_model is not None and y_model is not None:
            try:
                y_pred_obs = tobit.predict_observed(X_model)
                y_true_np = y_model.values

                # Predicho vs observado
                st.subheader("2. Predicho vs Observado")
                fig2 = plot_obs_vs_pred(y_true_np, y_pred_obs)
                st.pyplot(fig2, use_container_width=True)
                plt.close(fig2)

                divider()

                # Residuos y QQ
                st.subheader("3. Análisis de Residuos — QQ Plot")
                fig3 = plot_residuals(y_true_np, y_pred_obs)
                st.pyplot(fig3, use_container_width=True)
                plt.close(fig3)

                divider()

                # Forest plot
                marginal_df = st.session_state.get(
                    "marginal_df", compute_marginal_effects(tobit, X_model)
                )
                st.subheader("4. Forest Plot — Efectos Marginales Observados")
                fig4 = plot_forest(marginal_df)
                st.pyplot(fig4, use_container_width=True)
                plt.close(fig4)

                divider()

                # OLS vs Tobit
                st.subheader("5. Comparación OLS vs Tobit")
                fig5 = plot_ols_vs_tobit(X_model, y_model, tobit)
                st.pyplot(fig5, use_container_width=True)
                plt.close(fig5)

                divider()

                # Leverage / Influencia (Distancia de Cook aproximada)
                st.subheader("6. Influencia y Leverage (Distancia de Cook — Aproximación OLS)")
                try:
                    ols_inf = sm.OLS(y_model, X_model).fit()
                    influence = ols_inf.get_influence()
                    cook_d = influence.cooks_distance[0]
                    fig_cook, ax_cook = plt.subplots(figsize=(10, 4), dpi=150)
                    ax_cook.stem(range(len(cook_d)), cook_d, markerfmt="o",
                                 linefmt="gray", basefmt="black")
                    ax_cook.axhline(4 / len(cook_d), color="red", linestyle="--",
                                    linewidth=1, label=f"Umbral 4/n={4/len(cook_d):.4f}")
                    ax_cook.set_xlabel("Índice de Observación", fontsize=11)
                    ax_cook.set_ylabel("Distancia de Cook", fontsize=11)
                    ax_cook.set_title("Distancia de Cook — Detección de Observaciones Influyentes",
                                      fontsize=12, fontweight="bold", color=CAFE_OSCURO)
                    ax_cook.legend(fontsize=9)
                    ax_cook.grid(alpha=0.3)
                    st.pyplot(fig_cook, use_container_width=True)
                    plt.close(fig_cook)
                except Exception as exc_cook:
                    st.warning(f"No se pudo calcular distancia de Cook: {exc_cook}")

            except Exception as exc_vis:
                st.warning(
                    f"Ejecute el modelo en la pestaña **Modelado Tobit** para ver "
                    f"estas visualizaciones. ({exc_vis})"
                )
        else:
            st.info("🔬 Ejecute el modelo en la pestaña **Modelado Tobit** para ver "
                    "visualizaciones basadas en resultados.")

        divider()

        # Heatmap siempre disponible
        st.subheader("7. Matriz de Correlaciones")
        try:
            _, _, feat_cols_heat = preprocess(df, winsorize=False)
            n_vars = min(opts["max_vars_heatmap"], len(feat_cols_heat))
            top_corr_cols = (
                df[feat_cols_heat].corrwith(df[REQUIRED_COL])
                .abs().nlargest(n_vars).index.tolist()
            )
            df_heat = df[top_corr_cols + [REQUIRED_COL]].dropna()
            fig_hm2 = plot_heatmap(df_heat, top_corr_cols)
            st.pyplot(fig_hm2, use_container_width=True)
            plt.close(fig_hm2)
        except Exception as exc_hm:
            st.warning(f"Heatmap no disponible: {exc_hm}")

    # ============================================================
    # TAB 5 — EXPORTACIÓN
    # ============================================================
    with tabs[4]:
        st.header("📄 Exportación de Resultados")

        if df is None:
            st.info("⬆️ Cargue un dataset primero.")
            st.stop()

        tobit = st.session_state.get("tobit")
        if tobit is None:
            st.warning("⚠️ Ejecute el modelo en la pestaña **Modelado Tobit** antes de exportar.")
            st.stop()

        X_model = st.session_state.X
        y_model = st.session_state.y_ser
        feat_cols_exp = st.session_state.feat_cols
        fit_stats = st.session_state.fit_stats
        summary_df = st.session_state.get("summary_df", tobit.summary_df())
        marginal_df = st.session_state.get(
            "marginal_df", compute_marginal_effects(tobit, X_model)
        )

        try:
            _, df_model_exp, feat_cols_exp2 = preprocess(df, winsorize=opts["winsorize"])
        except Exception:
            df_model_exp = df.copy()
            feat_cols_exp2 = feat_cols_exp

        y_pred_exp = tobit.predict_observed(X_model)
        y_true_exp = y_model.values

        col_w, col_e = st.columns(2)

        with col_w:
            st.subheader("📝 Reporte Word (APA 7)")
            st.markdown(
                "Genera un documento `.docx` con portada, metodología, "
                "resultados, efectos marginales, conclusiones y referencias en "
                "formato APA 7 (Calibri 11, márgenes 1\")."
            )
            if st.button("📄 Generar Reporte Word", type="primary"):
                with st.spinner("Generando documento Word..."):
                    try:
                        word_bytes = generate_word_report(
                            df=df,
                            model=tobit,
                            summary=summary_df,
                            marginal=marginal_df,
                            fit_stats=fit_stats,
                            feature_cols=feat_cols_exp,
                        )
                        st.download_button(
                            label="⬇️ Descargar Reporte Word",
                            data=word_bytes,
                            file_name="Reporte_Tobit_PCS_Cenicafe.docx",
                            mime=(
                                "application/vnd.openxmlformats-officedocument"
                                ".wordprocessingml.document"
                            ),
                        )
                    except Exception as exc_w:
                        st.error(f"Error generando Word: {exc_w}")

        with col_e:
            st.subheader("📊 Reporte Excel (6 hojas)")
            st.markdown(
                "Genera un `.xlsx` con hojas: Datos_Modelo, Resultados_Tobit, "
                "Efectos_Marginales, Estadisticas, Predicciones, Residuos."
            )
            if st.button("📊 Generar Reporte Excel", type="primary"):
                with st.spinner("Generando Excel..."):
                    try:
                        excel_bytes = generate_excel_report(
                            df=df,
                            df_model=df_model_exp,
                            summary=summary_df,
                            marginal=marginal_df,
                            fit_stats=fit_stats,
                            y_pred=y_pred_exp,
                            y_true=y_true_exp,
                            feature_cols=feat_cols_exp,
                        )
                        st.download_button(
                            label="⬇️ Descargar Reporte Excel",
                            data=excel_bytes,
                            file_name="Reporte_Tobit_PCS_Cenicafe.xlsx",
                            mime=(
                                "application/vnd.openxmlformats-officedocument"
                                ".spreadsheetml.sheet"
                            ),
                        )
                    except Exception as exc_e:
                        st.error(f"Error generando Excel: {exc_e}")

    # ============================================================
    # TAB 6 — GITHUB & DEPLOY
    # ============================================================
    with tabs[5]:
        st.header("🚀 GitHub y Despliegue en Streamlit Cloud")

        st.subheader("📁 Estructura del Proyecto")
        st.code(
            """
📦 tobit-pcs-cenicafe/
├── app.py                          # Aplicación principal
├── requirements.txt                # Dependencias Python
├── TobitDataPCS_MASTER_Python.csv  # Dataset (si se incluye)
├── .gitignore                      # Archivos ignorados
└── README.md                       # Documentación
""",
            language="text",
        )

        st.subheader("🖥️ Terminal — Visual Studio Code")
        st.code(
            """
# 1. Crear entorno virtual
python -m venv venv

# 2. Activar entorno (Windows)
venv\\Scripts\\activate

# 3. Activar entorno (macOS/Linux)
source venv/bin/activate

# 4. Instalar dependencias
pip install -r requirements.txt

# 5. Ejecutar aplicación
streamlit run app.py
""",
            language="bash",
        )

        st.subheader("🐙 GitHub — Comandos")
        st.code(
            """
# Inicializar repositorio
git init

# Agregar archivos
git add .

# Commit inicial
git commit -m "feat: Modelo Tobit Cenicafé — PCS Huila 2026"

# Conectar con repositorio remoto
git remote add origin https://github.com/TU_USUARIO/tobit-pcs-cenicafe.git

# Subir código
git branch -M main
git push -u origin main
""",
            language="bash",
        )

        st.subheader("☁️ Despliegue en Streamlit Cloud")
        st.markdown(
            """
1. Ve a [share.streamlit.io](https://share.streamlit.io) e inicia sesión con GitHub.
2. Haz clic en **New app**.
3. Selecciona tu repositorio `tobit-pcs-cenicafe`.
4. Branch: `main` · Main file path: `app.py`.
5. Haz clic en **Deploy**.
6. En ~2 minutos tu app estará en:
   `https://TU_USUARIO-tobit-pcs-cenicafe-app.streamlit.app`
"""
        )

        st.subheader("📄 .gitignore recomendado")
        st.code(
            """
venv/
__pycache__/
*.pyc
.env
*.egg-info/
.DS_Store
Thumbs.db
""",
            language="text",
        )

        st.subheader("📋 README.md sugerido")
        st.code(
            """
# Modelo Tobit — Adopción de PCS | Cenicafé 2026

Aplicación Streamlit para análisis econométrico de intensidad de adopción
de Prácticas de Conservación de Suelos (PCS) en fincas cafeteras del Huila.

## Modelo
Tobit de doble censura (0-100) estimado por Máxima Verosimilitud.

## Dataset
422 observaciones · 5 municipios del Huila · Variable dependiente: Intensidad

## Instalación
```bash
pip install -r requirements.txt
streamlit run app.py
```

## Referencia
Tobin, J. (1958). Estimation of relationships for limited dependent variables.
*Econometrica*, 26(1), 24–36.
""",
            language="markdown",
        )

        divider()
        st.markdown(
            f"<p style='color:{CAFE_OSCURO};font-size:13px;'>"
            "☕ <b>Cenicafé</b> · Centro Nacional de Investigaciones de Café · "
            "Chinchiná, Caldas, Colombia · 2026"
            "</p>",
            unsafe_allow_html=True,
        )


if __name__ == "__main__":
    main()
